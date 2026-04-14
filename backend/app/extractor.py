"""Extract plain text from PDF and DOCX files."""
import io
import re


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the right extractor based on file extension.

    Returns clean plain text.  Raises ``ValueError`` for unsupported types
    or when the file yields no extractable text.
    """
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    if ext == "pdf":
        text = _extract_pdf(file_bytes)
    elif ext == "docx":
        text = _extract_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: .{ext}. Upload a .pdf or .docx file."
        )

    text = _clean(text)

    if not text:
        raise ValueError(
            "No text could be extracted from this file. "
            "It may be a scanned image — OCR is not supported yet."
        )

    return text


# ---------------------------------------------------------------------------
# PDF  (PyMuPDF / fitz)
# ---------------------------------------------------------------------------

def _extract_pdf(data: bytes) -> str:
    import fitz  # pymupdf

    pages: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            pages.append(page.get_text("text"))
    return "\n\n".join(pages)


# ---------------------------------------------------------------------------
# DOCX  (python-docx)
# ---------------------------------------------------------------------------

def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))

    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables (row by row, cells tab-separated)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Cleanup
# ---------------------------------------------------------------------------

def _clean(text: str) -> str:
    # Collapse runs of 3+ newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip trailing whitespace on each line
    text = "\n".join(ln.rstrip() for ln in text.splitlines())
    return text.strip()
